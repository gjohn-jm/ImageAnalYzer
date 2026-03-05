import os
from pathlib import Path
import base64
import mimetypes
import io
import json

from PIL import Image, ImageOps, ImageEnhance
import pytesseract
from openpyxl import Workbook

from groq import Groq
from docx import Document
from fpdf import FPDF

from PySide6.QtCore import Qt, QRect, Signal, QTimer, QStandardPaths
from PySide6.QtGui import QPixmap, QPainter, QPen, QFont, QAction, QTextOption,QIcon
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QPushButton, QFileDialog,
    QPlainTextEdit, QHBoxLayout, QVBoxLayout, QGroupBox, QMessageBox,
    QLineEdit, QComboBox, QCheckBox, QTabWidget, QStatusBar, QToolBar,
    QSplitter, QSizePolicy, QScrollArea, QInputDialog
)

TESSERACT_EXE = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


class SelectableImageLabel(QLabel):
    selectionFinished = Signal()

    def __init__(self):
        super().__init__()
        self.setMouseTracking(True)
        self.selection = QRect()
        self._dragging = False
        self._start = None

    def clear_selection(self):
        self.selection = QRect()
        self._dragging = False
        self._start = None
        self.update()

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._dragging = True
            self._start = event.position().toPoint()
            self.selection = QRect(self._start, self._start)
            self.update()

    def mouseMoveEvent(self, event):
        if self._dragging and self._start is not None:
            cur = event.position().toPoint()
            self.selection = QRect(self._start, cur).normalized()
            self.update()

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton and self._dragging:
            self._dragging = False
            cur = event.position().toPoint()
            self.selection = QRect(self._start, cur).normalized()
            self.update()
            self.selectionFinished.emit()

    def paintEvent(self, event):
        super().paintEvent(event)
        if not self.selection.isNull() and self.selection.width() > 6 and self.selection.height() > 6:
            painter = QPainter(self)
            painter.setRenderHint(QPainter.Antialiasing)
            painter.setPen(QPen(Qt.green, 2))
            painter.drawRect(self.selection)


class Image2ExcelGroqApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ImageAnalYzer")
        self.setWindowIcon(QIcon(r"C:\Users\raidenxr\Downloads\Analyzer.ico"))
        self.resize(1520, 880)

        self.image_path = None
        self._pixmap_original = None
        self._restoring_splitters = False

        self._lens_zoom = 1.0
        self._lens_zoom_min = 0.4
        self._lens_zoom_max = 4.0
        self._lens_zoom_step = 1.25

        self._build_ui()
        self._connect()

        self.set_status("Ready. Open an image.")

        if not self._load_api_key() and not (os.environ.get("GROQ_API_KEY") or os.environ.get("GROQAPIKEY")):
            self.statusBar().showMessage("Tip: Click 'Set API Key' to add your Groq key.")

    # ---------------- API key storage ----------------
    def _config_dir(self) -> Path:
        base = Path(QStandardPaths.writableLocation(QStandardPaths.AppDataLocation))
        cfg = base / "Image2ExcelAI"
        cfg.mkdir(parents=True, exist_ok=True)
        return cfg

    def _config_path(self) -> Path:
        return self._config_dir() / "config.json"

    def _load_api_key(self) -> str:
        try:
            p = self._config_path()
            if not p.exists():
                return ""
            data = json.loads(p.read_text(encoding="utf-8"))
            return (data.get("groq_api_key") or "").strip()
        except Exception:
            return ""

    def _save_api_key(self, key: str) -> None:
        p = self._config_path()
        data = {"groq_api_key": (key or "").strip()}
        p.write_text(json.dumps(data, indent=2), encoding="utf-8")

    def set_api_key_dialog(self):
        text, ok = QInputDialog.getText(
            self,
            "Set Groq API Key",
            "Paste your Groq API key (saved for this Windows user):",
            QLineEdit.Password,
            ""
        )
        if not ok:
            return

        key = (text or "").strip()
        if not key:
            QMessageBox.warning(self, "Empty key", "API key not saved (empty).")
            return

        self._save_api_key(key)
        QMessageBox.information(self, "Saved", "API key saved. You can now use AI features.")
        self.set_status("API key saved.")

    # ---------------- UI helpers ----------------
    def _setup_plaintext(self, w: QPlainTextEdit, font_size=12, read_only=False, always_scroll=False):
        w.setReadOnly(read_only)
        w.setFont(QFont("Consolas", font_size))

        # Prevent “text goes outside” for long words/no-spaces
        w.setLineWrapMode(QPlainTextEdit.WidgetWidth)
        w.setWordWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)

        # Force no horizontal overflow; rely on wrapping
        w.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        # For the context box we force a visible scrollbar
        if always_scroll:
            w.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOn)
        else:
            w.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)

        w.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

    def _setup_lineedit(self, w: QLineEdit):
        w.setClearButtonEnabled(True)
        w.setMinimumHeight(36)

    # ---------------- UI ----------------
    def _build_ui(self):
        # Toolbar
        toolbar = QToolBar("Main")
        toolbar.setMovable(False)
        self.addToolBar(toolbar)

        self.act_open = QAction("Open", self)
        self.act_remove_image = QAction("Remove Image", self)
        self.act_extract = QAction("Extract OCR", self)
        self.act_export_excel = QAction("Export Excel", self)
        self.act_export_word = QAction("Export Word", self)
        self.act_export_pdf = QAction("Export PDF", self)
        self.act_clear = QAction("Clear Selection", self)
        self.act_set_key = QAction("Set API Key", self)

        toolbar.addAction(self.act_open)
        toolbar.addAction(self.act_remove_image)
        toolbar.addSeparator()
        toolbar.addAction(self.act_extract)
        toolbar.addAction(self.act_export_excel)
        toolbar.addAction(self.act_export_word)
        toolbar.addAction(self.act_export_pdf)
        toolbar.addSeparator()
        toolbar.addAction(self.act_clear)
        toolbar.addSeparator()
        toolbar.addAction(self.act_set_key)

        # Status bar
        self.setStatusBar(QStatusBar(self))
        self.statusBar().showMessage("Ready")

        # ---------- Preview (tab 1) ----------
        self.image_label_ocr = QLabel("Open Image")
        self.image_label_ocr.setAlignment(Qt.AlignCenter)
        self.image_label_ocr.setMinimumHeight(420)
        self.image_label_ocr.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Ignored)

        self.preview_box_ocr = QGroupBox("Preview")
        preview_ocr_l = QVBoxLayout()
        preview_ocr_l.addWidget(self.image_label_ocr)
        self.preview_box_ocr.setLayout(preview_ocr_l)
        self.preview_box_ocr.setMinimumWidth(380)
        self.preview_box_ocr.setMaximumWidth(460)

        # ---------- Preview (tab 2, selectable + zoomable) ----------
        self.image_label = SelectableImageLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setText("Open Image → Drag to select area")
        self.image_label.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        self.image_label.setMinimumSize(500, 360)

        self.lens_scroll = QScrollArea()
        self.lens_scroll.setWidget(self.image_label)
        self.lens_scroll.setWidgetResizable(False)
        self.lens_scroll.setAlignment(Qt.AlignCenter)
        self.lens_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.lens_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)

        self.btn_zoom_out = QPushButton("−")
        self.btn_zoom_in = QPushButton("+")
        self.btn_zoom_reset = QPushButton("Reset")
        self.lbl_zoom = QLabel("Zoom: 100%")
        self.lbl_zoom.setMinimumWidth(110)

        zoom_row = QWidget()
        zoom_l = QHBoxLayout()
        zoom_l.setContentsMargins(0, 0, 0, 0)
        zoom_l.setSpacing(8)
        zoom_l.addWidget(self.btn_zoom_out)
        zoom_l.addWidget(self.btn_zoom_in)
        zoom_l.addWidget(self.btn_zoom_reset)
        zoom_l.addStretch(1)
        zoom_l.addWidget(self.lbl_zoom)
        zoom_row.setLayout(zoom_l)

        self.preview_box_lens = QGroupBox("Preview (Zoom + Select)")
        preview_lens_l = QVBoxLayout()
        preview_lens_l.setSpacing(8)
        preview_lens_l.addWidget(self.lens_scroll, 1)
        preview_lens_l.addWidget(zoom_row)
        self.preview_box_lens.setLayout(preview_lens_l)
        self.preview_box_lens.setMinimumWidth(340)
        self.preview_box_lens.setMaximumWidth(520)

        # ---------- Tab 1: OCR & Export ----------
        self.full_text = QPlainTextEdit()
        self._setup_plaintext(self.full_text, font_size=12, read_only=False)
        self.full_text.setPlaceholderText("Full-image extracted text will appear here...")

        self.btn_improve_full = QPushButton("Increase Accuracy (Full Text)")

        export_row = QWidget()
        export_layout = QHBoxLayout()
        export_layout.setContentsMargins(0, 0, 0, 0)
        export_layout.setSpacing(10)

        self.btn_export_excel = QPushButton("Export Excel")
        self.btn_export_word = QPushButton("Export Word")
        self.btn_export_pdf = QPushButton("Export PDF")

        export_layout.addWidget(self.btn_export_excel)
        export_layout.addWidget(self.btn_export_word)
        export_layout.addWidget(self.btn_export_pdf)
        export_row.setLayout(export_layout)

        full_box = QGroupBox("Full OCR Text (editable)")
        full_l = QVBoxLayout()
        full_l.addWidget(self.full_text, 1)
        full_l.addWidget(self.btn_improve_full)
        full_l.addWidget(export_row)
        full_box.setLayout(full_l)

        self.tab1_split = QSplitter(Qt.Horizontal)
        self.tab1_split.setChildrenCollapsible(False)
        self.tab1_split.addWidget(self.preview_box_ocr)
        self.tab1_split.addWidget(full_box)
        self.tab1_split.setStretchFactor(0, 1)
        self.tab1_split.setStretchFactor(1, 5)

        tab1 = QWidget()
        tab1_l = QHBoxLayout()
        tab1_l.setContentsMargins(12, 12, 12, 12)
        tab1_l.setSpacing(12)
        tab1_l.addWidget(self.tab1_split, 1)
        tab1.setLayout(tab1_l)

        # ---------- Tab 2: Lens OCR ----------
        self.selection_text = QPlainTextEdit()
        self._setup_plaintext(self.selection_text, font_size=12, read_only=True)
        self.selection_text.setPlaceholderText("Selected-area OCR will appear here after you drag.")
        self.selection_text.setMinimumHeight(200)

        self.improved_text = QPlainTextEdit()
        self._setup_plaintext(self.improved_text, font_size=12, read_only=False)
        self.improved_text.setPlaceholderText("Improved (fixed) text will appear here...")
        self.improved_text.setMinimumHeight(240)

        self.model_box = QComboBox()
        self.model_box.addItems([
            "llama-3.3-70b-versatile",
            "llama-3.1-8b-instant",
            "mixtral-8x7b-32768",
        ])
        self.model_box.setCurrentText("llama-3.3-70b-versatile")

        self.auto_answer = QCheckBox("Auto-answer after selection (uses AI Tools tab)")
        self.chk_use_vision_for_accuracy = QCheckBox("Use Vision re-read for Increase Accuracy (best for blur)")

        self.btn_improve_sel = QPushButton("Increase Accuracy (Selection)")
        self.btn_send_sel_to_ai = QPushButton("Send Selected OCR → AI Tools")
        self.btn_send_improved_to_ai = QPushButton("Send Improved Text → AI Tools")

        lens_box = QGroupBox("Selection OCR")
        lens_l = QVBoxLayout()
        lens_l.setSpacing(10)
        lens_l.addWidget(QLabel("Selected-area OCR:"))
        lens_l.addWidget(self.selection_text, 1)
        lens_l.addWidget(QLabel("Improved text:"))
        lens_l.addWidget(self.improved_text, 1)

        lens_l.addWidget(QLabel("Text model:"))
        lens_l.addWidget(self.model_box)
        lens_l.addWidget(self.chk_use_vision_for_accuracy)
        lens_l.addWidget(self.auto_answer)

        row_btns = QWidget()
        row_btns_l = QHBoxLayout()
        row_btns_l.setContentsMargins(0, 0, 0, 0)
        row_btns_l.setSpacing(10)
        row_btns_l.addWidget(self.btn_improve_sel)
        row_btns_l.addWidget(self.btn_send_sel_to_ai)
        row_btns_l.addWidget(self.btn_send_improved_to_ai)
        row_btns.setLayout(row_btns_l)
        lens_l.addWidget(row_btns)

        lens_box.setLayout(lens_l)

        self.tab2_split = QSplitter(Qt.Horizontal)
        self.tab2_split.setChildrenCollapsible(False)
        self.tab2_split.addWidget(self.preview_box_lens)
        self.tab2_split.addWidget(lens_box)
        self.tab2_split.setStretchFactor(0, 1)
        self.tab2_split.setStretchFactor(1, 5)

        tab2 = QWidget()
        tab2_l = QHBoxLayout()
        tab2_l.setContentsMargins(12, 12, 12, 12)
        tab2_l.setSpacing(12)
        tab2_l.addWidget(self.tab2_split, 1)
        tab2.setLayout(tab2_l)

        # ---------- Tab 3: AI Tools ----------
        self.ai_context = QPlainTextEdit()
        # Force visible vertical scrollbar + wrap so text never draws outside
        self._setup_plaintext(self.ai_context, font_size=12, read_only=False, always_scroll=True)
        self.ai_context.setPlaceholderText("Context for AI (paste here or use the Send buttons)...")
        self.ai_context.setMinimumHeight(260)

        self.ai_question = QLineEdit()
        self._setup_lineedit(self.ai_question)
        self.ai_question.setPlaceholderText("Ask a question about the context...")

        self.btn_ask_ai = QPushButton("Ask AI")
        self.btn_ai_fix_context = QPushButton("Increase Accuracy (Fix Context Text)")

        self.btn_analyze_image = QPushButton("Analyze Image (Vision)")
        self.btn_vision_reread = QPushButton("Vision Re-read → Context (best for blur)")
        self.chk_vision_use_selection = QCheckBox("Vision uses selection crop (if available)")
        self.chk_vision_use_selection.setChecked(True)

        self.ai_answer = QPlainTextEdit()
        self._setup_plaintext(self.ai_answer, font_size=12, read_only=True)
        self.ai_answer.setPlaceholderText("AI answer will appear here...")
        self.ai_answer.setMinimumHeight(420)

        self.image_analysis = QPlainTextEdit()
        self._setup_plaintext(self.image_analysis, font_size=12, read_only=True)
        self.image_analysis.setPlaceholderText("Vision analysis will appear here...")
        self.image_analysis.setMinimumHeight(220)

        # Build AI Inputs content widget, then wrap it in a scroll area
        ai_inputs_widget = QWidget()
        ai_inputs_l = QVBoxLayout()
        ai_inputs_l.setSpacing(10)
        ai_inputs_l.setContentsMargins(10, 12, 10, 12)

        ctx_btns = QWidget()
        ctx_btns_l = QHBoxLayout()
        ctx_btns_l.setContentsMargins(0, 0, 0, 0)
        ctx_btns_l.setSpacing(10)
        self.btn_use_full_ocr_to_ai = QPushButton("Use Full OCR")
        self.btn_use_sel_ocr_to_ai = QPushButton("Use Selected OCR")
        self.btn_use_improved_to_ai = QPushButton("Use Improved Text")
        ctx_btns_l.addWidget(self.btn_use_full_ocr_to_ai)
        ctx_btns_l.addWidget(self.btn_use_sel_ocr_to_ai)
        ctx_btns_l.addWidget(self.btn_use_improved_to_ai)
        ctx_btns.setLayout(ctx_btns_l)

        ai_inputs_l.addWidget(QLabel("Context:"))
        ai_inputs_l.addWidget(self.ai_context, 1)
        ai_inputs_l.addWidget(ctx_btns)

        q_row = QWidget()
        q_row_l = QHBoxLayout()
        q_row_l.setContentsMargins(0, 0, 0, 0)
        q_row_l.setSpacing(10)
        q_row_l.addWidget(self.ai_question, 1)
        q_row_l.addWidget(self.btn_ask_ai)
        q_row.setLayout(q_row_l)
        ai_inputs_l.addWidget(q_row)

        tools_row = QWidget()
        tools_row_l = QHBoxLayout()
        tools_row_l.setContentsMargins(0, 0, 0, 0)
        tools_row_l.setSpacing(10)
        tools_row_l.addWidget(self.btn_ai_fix_context)
        tools_row_l.addWidget(self.btn_vision_reread)
        tools_row_l.addWidget(self.btn_analyze_image)
        tools_row.setLayout(tools_row_l)

        ai_inputs_l.addWidget(self.chk_vision_use_selection)
        ai_inputs_l.addWidget(tools_row)
        ai_inputs_widget.setLayout(ai_inputs_l)

        top_controls = QGroupBox("AI Inputs")
        top_controls_l = QVBoxLayout()
        top_controls_l.setContentsMargins(0, 0, 0, 0)

        # This scrolls the entire AI Inputs panel if space is tight
        top_scroll = QScrollArea()
        top_scroll.setWidget(ai_inputs_widget)
        top_scroll.setWidgetResizable(True)
        top_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        top_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)

        top_controls_l.addWidget(top_scroll, 1)
        top_controls.setLayout(top_controls_l)

        outputs_tabs = QTabWidget()
        out_answer = QWidget()
        out_answer_l = QVBoxLayout()
        out_answer_l.setContentsMargins(0, 0, 0, 0)
        out_answer_l.addWidget(self.ai_answer, 1)
        out_answer.setLayout(out_answer_l)

        out_vision = QWidget()
        out_vision_l = QVBoxLayout()
        out_vision_l.setContentsMargins(0, 0, 0, 0)
        out_vision_l.addWidget(self.image_analysis, 1)
        out_vision.setLayout(out_vision_l)

        outputs_tabs.addTab(out_answer, "Answer")
        outputs_tabs.addTab(out_vision, "Vision output")

        self.tab3_split = QSplitter(Qt.Vertical)
        self.tab3_split.setChildrenCollapsible(False)
        self.tab3_split.addWidget(top_controls)
        self.tab3_split.addWidget(outputs_tabs)

        tab3 = QWidget()
        tab3_l = QVBoxLayout()
        tab3_l.setContentsMargins(12, 12, 12, 12)
        tab3_l.setSpacing(12)
        tab3_l.addWidget(self.tab3_split, 1)
        tab3.setLayout(tab3_l)

        # Tabs
        self.tabs = QTabWidget()
        self.tabs.addTab(tab1, "OCR & Export")
        self.tabs.addTab(tab2, "Lens OCR")
        self.tabs.addTab(tab3, "AI Tools")

        root = QWidget()
        root_l = QVBoxLayout()
        root_l.setContentsMargins(10, 10, 10, 10)
        root_l.addWidget(self.tabs, 1)
        root.setLayout(root_l)
        self.setCentralWidget(root)

        QTimer.singleShot(0, self._restore_splitter_defaults)

    # ---------------- Zoom (Lens) ----------------
    def _set_lens_zoom(self, value: float):
        value = max(self._lens_zoom_min, min(self._lens_zoom_max, float(value)))
        if abs(value - self._lens_zoom) < 1e-6:
            return
        self._lens_zoom = value
        self.lbl_zoom.setText(f"Zoom: {int(round(self._lens_zoom * 100))}%")
        self.image_label.clear_selection()
        self._update_lens_pixmap()

    def zoom_in(self):
        self._set_lens_zoom(self._lens_zoom * self._lens_zoom_step)

    def zoom_out(self):
        self._set_lens_zoom(self._lens_zoom / self._lens_zoom_step)

    def zoom_reset(self):
        self._set_lens_zoom(1.0)

    def _update_lens_pixmap(self):
        if not self._pixmap_original:
            self.image_label.setPixmap(QPixmap())
            self.image_label.setText("Open Image → Drag to select area")
            self.image_label.setFixedSize(
                max(500, self.lens_scroll.viewport().width()),
                max(360, self.lens_scroll.viewport().height())
            )
            return

        vw = max(100, self.lens_scroll.viewport().width())
        vh = max(100, self.lens_scroll.viewport().height())

        pm_fit = self._pixmap_original.scaled(vw, vh, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        target_w = max(1, int(pm_fit.width() * self._lens_zoom))
        target_h = max(1, int(pm_fit.height() * self._lens_zoom))

        pm_zoom = self._pixmap_original.scaled(target_w, target_h, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.image_label.setText("")
        self.image_label.setPixmap(pm_zoom)
        self.image_label.setFixedSize(pm_zoom.size())

    # ---------------- Splitter stability ----------------
    def _restore_splitter_defaults(self):
        if self._restoring_splitters:
            return
        self._restoring_splitters = True
        try:
            self.tab1_split.setSizes([420, 1200])
            self.tab2_split.setSizes([480, 1140])
            self.tab3_split.setSizes([360, 520])
        finally:
            self._restoring_splitters = False

    def _restore_splitters_soon(self):
        QTimer.singleShot(0, self._restore_splitter_defaults)

    # ---------------- Helpers ----------------
    def set_status(self, msg: str):
        self.statusBar().showMessage(msg)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._refresh_preview()

    def _refresh_preview(self):
        if not self._pixmap_original:
            self.image_label_ocr.setPixmap(QPixmap())
            self.image_label_ocr.setText("Open Image")
        else:
            w1, h1 = self.image_label_ocr.width(), self.image_label_ocr.height()
            pm1 = self._pixmap_original.scaled(w1, h1, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.image_label_ocr.setText("")
            self.image_label_ocr.setPixmap(pm1)

        self._update_lens_pixmap()

    def _ensure_tesseract(self):
        if not os.path.exists(TESSERACT_EXE):
            QMessageBox.critical(self, "Tesseract not found", f"Not found:\n{TESSERACT_EXE}\nEdit TESSERACT_EXE.")
            return False
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_EXE
        return True

    def _ensure_groq(self):
        key = self._load_api_key()
        if not key:
            key = (os.environ.get("GROQ_API_KEY") or os.environ.get("GROQAPIKEY") or "").strip()

        if not key:
            QMessageBox.critical(
                self,
                "Missing Groq API key",
                "Click 'Set API Key' in the toolbar and paste your key.\n\n"
                "Or set GROQ_API_KEY in Environment Variables."
            )
            return None

        return Groq(api_key=key)

    def _ocr_preprocess_pil(self, img: Image.Image) -> Image.Image:
        img = ImageOps.exif_transpose(img)
        img = img.convert("L")
        img = ImageEnhance.Contrast(img).enhance(1.8)
        img = ImageEnhance.Sharpness(img).enhance(1.5)
        img = img.point(lambda p: 255 if p > 170 else 0)
        return img

    def _ocr_pil(self, img: Image.Image) -> str:
        if not self._ensure_tesseract():
            return ""
        cfg = r"--oem 3 --psm 6"
        img2 = self._ocr_preprocess_pil(img)
        return (pytesseract.image_to_string(img2, lang="eng", config=cfg) or "").strip()

    def _selection_crop_box_original_pixels(self):
        if not self._pixmap_original:
            return None

        pm = self.image_label.pixmap()
        if pm is None:
            return None

        sel = self.image_label.selection
        if sel.isNull() or sel.width() < 10 or sel.height() < 10:
            return None

        disp_w, disp_h = pm.width(), pm.height()
        lab_w, lab_h = self.image_label.width(), self.image_label.height()

        off_x = (lab_w - disp_w) // 2
        off_y = (lab_h - disp_h) // 2

        x1 = sel.left() - off_x
        y1 = sel.top() - off_y
        x2 = sel.right() - off_x
        y2 = sel.bottom() - off_y

        x1 = max(0, min(x1, disp_w - 1))
        y1 = max(0, min(y1, disp_h - 1))
        x2 = max(0, min(x2, disp_w - 1))
        y2 = max(0, min(y2, disp_h - 1))

        if x2 <= x1 or y2 <= y1:
            return None

        orig_w, orig_h = self._pixmap_original.width(), self._pixmap_original.height()
        fx = orig_w / disp_w
        fy = orig_h / disp_h

        left = int(x1 * fx)
        top = int(y1 * fy)
        right = int((x2 + 1) * fx)
        bottom = int((y2 + 1) * fy)

        left = max(0, min(left, orig_w - 1))
        top = max(0, min(top, orig_h - 1))
        right = max(left + 1, min(right, orig_w))
        bottom = max(top + 1, min(bottom, orig_h))

        return (left, top, right, bottom)

    def _image_to_data_url(self, pil_img: Image.Image, fallback_ext: str = ".jpg") -> str:
        ext = (Path(self.image_path).suffix.lower() if self.image_path else fallback_ext).lower()
        mime = mimetypes.types_map.get(ext, "image/jpeg")

        buf = io.BytesIO()
        if mime == "image/png":
            pil_img.save(buf, format="PNG")
        else:
            pil_img = pil_img.convert("RGB")
            pil_img.save(buf, format="JPEG", quality=92)

        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
        return f"data:{mime};base64,{b64}"

    # ---------------- Actions ----------------
    def _connect(self):
        self.act_open.triggered.connect(self.select_image)
        self.act_remove_image.triggered.connect(self.remove_image)
        self.act_extract.triggered.connect(self.extract_full_text)
        self.act_export_excel.triggered.connect(self.export_excel)
        self.act_export_word.triggered.connect(self.export_word)
        self.act_export_pdf.triggered.connect(self.export_pdf)
        self.act_clear.triggered.connect(self.clear_selection)
        self.act_set_key.triggered.connect(self.set_api_key_dialog)

        self.image_label.selectionFinished.connect(self.auto_ocr_selection)

        self.btn_export_excel.clicked.connect(self.export_excel)
        self.btn_export_word.clicked.connect(self.export_word)
        self.btn_export_pdf.clicked.connect(self.export_pdf)

        self.btn_improve_full.clicked.connect(self.increase_accuracy_full)
        self.btn_improve_sel.clicked.connect(self.increase_accuracy_selection)

        self.btn_send_sel_to_ai.clicked.connect(self.send_selected_to_ai_tab)
        self.btn_send_improved_to_ai.clicked.connect(self.send_improved_to_ai_tab)

        self.btn_use_full_ocr_to_ai.clicked.connect(self.use_full_ocr_in_ai)
        self.btn_use_sel_ocr_to_ai.clicked.connect(self.use_selected_ocr_in_ai)
        self.btn_use_improved_to_ai.clicked.connect(self.use_improved_in_ai)

        self.btn_ask_ai.clicked.connect(self.ask_ai_from_ai_tab)
        self.ai_question.returnPressed.connect(self.ask_ai_from_ai_tab)

        self.btn_ai_fix_context.clicked.connect(self.increase_accuracy_ai_context)
        self.btn_analyze_image.clicked.connect(self.analyze_image_vision)
        self.btn_vision_reread.clicked.connect(self.vision_reread_to_ai_context)

        self.btn_zoom_in.clicked.connect(self.zoom_in)
        self.btn_zoom_out.clicked.connect(self.zoom_out)
        self.btn_zoom_reset.clicked.connect(self.zoom_reset)

    # ---------- your existing methods below (select_image/remove_image/etc) ----------
    # NOTE: Everything else stays the same as your earlier version.

    def select_image(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Select an image", str(Path.home()),
            "Images (*.png *.jpg *.jpeg *.bmp *.tif *.tiff)"
        )
        if not path:
            return

        pix = QPixmap(path)
        if pix.isNull():
            QMessageBox.critical(self, "Error", "Could not load the image.")
            return

        self.image_path = path
        self._pixmap_original = pix

        self._lens_zoom = 1.0
        self.lbl_zoom.setText("Zoom: 100%")

        self.image_label.clear_selection()
        self.full_text.setPlainText("")
        self.selection_text.setPlainText("")
        self.improved_text.setPlainText("")
        self.ai_context.setPlainText("")
        self.ai_answer.setPlainText("")
        self.ai_question.setText("")
        self.image_analysis.setPlainText("")

        self.set_status(f"Selected: {Path(path).name}")
        self._refresh_preview()
        self._restore_splitters_soon()

    def remove_image(self):
        if not self.image_path:
            return

        self.image_path = None
        self._pixmap_original = None

        self._lens_zoom = 1.0
        self.lbl_zoom.setText("Zoom: 100%")

        self.image_label.clear_selection()
        self.full_text.setPlainText("")
        self.selection_text.setPlainText("")
        self.improved_text.setPlainText("")
        self.ai_context.setPlainText("")
        self.ai_answer.setPlainText("")
        self.ai_question.setText("")
        self.image_analysis.setPlainText("")

        self._refresh_preview()
        self.set_status("Image removed.")
        self._restore_splitters_soon()

    def clear_selection(self):
        self.image_label.clear_selection()
        self.selection_text.setPlainText("")
        self.improved_text.setPlainText("")
        self.set_status("Selection cleared.")

    def extract_full_text(self):
        if not self.image_path:
            QMessageBox.warning(self, "Missing image", "Select an image first.")
            return
        if not self._ensure_tesseract():
            return

        try:
            self.set_status("Extracting full text...")
            img = Image.open(self.image_path)
            text = self._ocr_pil(img)
            self.full_text.setPlainText(text)
            self.set_status("Full text extracted.")
        except Exception as e:
            QMessageBox.critical(self, "OCR error", str(e))
            self.set_status("Full OCR failed.")

    def export_excel(self):
        text = self.full_text.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "No text", "Extract full text first (or paste text), then Export.")
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self, "Save Excel file", str(Path.home() / "output.xlsx"),
            "Excel Workbook (*.xlsx)"
        )
        if not save_path:
            return
        if not save_path.lower().endswith(".xlsx"):
            save_path += ".xlsx"

        try:
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            wb = Workbook()
            ws = wb.active
            ws.title = "OCR"
            for i, line in enumerate(lines, start=1):
                ws.cell(row=i, column=1, value=line)
            wb.save(save_path)
            self.set_status(f"Excel saved: {Path(save_path).name}")
            QMessageBox.information(self, "Excel Exported", f"Saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Export error", str(e))
            self.set_status("Export failed.")

    def export_word(self):
        text = self.full_text.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "No text", "Extract full text first (or paste text), then Export.")
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self, "Save Word document", str(Path.home() / "output.docx"),
            "Word Document (*.docx)"
        )
        if not save_path:
            return
        if not save_path.lower().endswith(".docx"):
            save_path += ".docx"

        try:
            self.set_status("Creating Word document...")
            doc = Document()
            doc.add_heading("OCR Extracted Text", 0)
            for para in text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.save(save_path)
            self.set_status(f"Word saved: {Path(save_path).name}")
            QMessageBox.information(self, "Word Exported", f"Saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Word Export error", str(e))
            self.set_status("Word export failed.")

    def export_pdf(self):
        text = self.full_text.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "No text", "Extract full text first (or paste text), then Export.")
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF document", str(Path.home() / "output.pdf"),
            "PDF Document (*.pdf)"
        )
        if not save_path:
            return
        if not save_path.lower().endswith(".pdf"):
            save_path += ".pdf"

        try:
            self.set_status("Creating PDF document...")
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.set_font("Arial", style="B", size=16)
            pdf.cell(0, 10, txt="OCR Extracted Text", ln=True, align="C")
            pdf.ln(5)

            pdf.set_font("Arial", size=11)
            for line in text.split("\n"):
                try:
                    safe_line = line.encode("latin-1", "replace").decode("latin-1")
                except Exception:
                    safe_line = "".join(c if ord(c) < 128 else "?" for c in line)
                pdf.multi_cell(0, 6, txt=safe_line)

            pdf.output(save_path)
            self.set_status(f"PDF saved: {Path(save_path).name}")
            QMessageBox.information(self, "PDF Exported", f"Saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "PDF Export error", str(e))
            self.set_status("PDF export failed.")

    def auto_ocr_selection(self):
        if not self.image_path:
            return

        box = self._selection_crop_box_original_pixels()
        if not box:
            self.selection_text.setPlainText("")
            return

        try:
            self.set_status("Auto OCR selection...")
            img = Image.open(self.image_path)
            crop = img.crop(box)
            text = self._ocr_pil(crop)
            self.selection_text.setPlainText(text)
            self.improved_text.setPlainText("")

            if self.auto_answer.isChecked():
                self.ai_context.setPlainText(text)
                self.tabs.setCurrentIndex(2)
                self.ai_question.setText("Explain the context in detail (steps, key points, numbers).")
                self.ask_ai_from_ai_tab()
            else:
                self.set_status("Selection text detected.")
        except Exception as e:
            self.selection_text.setPlainText("")
            self.set_status("Selection OCR failed.")
            QMessageBox.critical(self, "Selection OCR error", str(e))

    def send_selected_to_ai_tab(self):
        txt = self.selection_text.toPlainText().strip()
        if not txt:
            QMessageBox.warning(self, "No selection text", "Select an area first.")
            return
        self.ai_context.setPlainText(txt)
        self.tabs.setCurrentIndex(2)

    def send_improved_to_ai_tab(self):
        txt = self.improved_text.toPlainText().strip()
        if not txt:
            QMessageBox.warning(self, "No improved text", "Generate improved text first.")
            return
        self.ai_context.setPlainText(txt)
        self.tabs.setCurrentIndex(2)

    def use_full_ocr_in_ai(self):
        txt = self.full_text.toPlainText().strip()
        if not txt:
            QMessageBox.warning(self, "No full OCR", "Extract full OCR first.")
            return
        self.ai_context.setPlainText(txt)

    def use_selected_ocr_in_ai(self):
        txt = self.selection_text.toPlainText().strip()
        if not txt:
            QMessageBox.warning(self, "No selection OCR", "Select an area first.")
            return
        self.ai_context.setPlainText(txt)

    def use_improved_in_ai(self):
        txt = self.improved_text.toPlainText().strip()
        if not txt:
            QMessageBox.warning(self, "No improved text", "Generate improved text first.")
            return
        self.ai_context.setPlainText(txt)

    def ask_ai_from_ai_tab(self):
        context = self.ai_context.toPlainText().strip()
        if not context:
            QMessageBox.warning(self, "No context", "Put some text in the AI context box first.")
            return

        q = self.ai_question.text().strip()
        if not q:
            QMessageBox.warning(self, "No question", "Type a question first.")
            return

        client = self._ensure_groq()
        if client is None:
            return

        try:
            self.set_status("Asking Groq...")
            model = self.model_box.currentText()

            completion = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a detailed assistant. Use ONLY the provided context text. "
                            "If the answer is not in the context, say: Not found in the context."
                        )
                    },
                    {"role": "user", "content": f"CONTEXT:\n{context}\n\nQUESTION:\n{q}"}
                ],
                temperature=0.4,
                max_tokens=900,
            )

            answer = completion.choices[0].message.content
            self.ai_answer.setPlainText((answer or "").strip())
            self.set_status("Answered.")
        except Exception as e:
            self.set_status("AI failed.")
            QMessageBox.critical(self, "Groq error", str(e))

    def increase_accuracy_full(self):
        raw = self.full_text.toPlainText().strip()
        if not raw:
            QMessageBox.warning(self, "No OCR text", "Extract full OCR first, then click Increase Accuracy.")
            return
        fixed = self._ai_fix_ocr_text(raw)
        if fixed is not None:
            self.full_text.setPlainText(fixed)

    def increase_accuracy_selection(self):
        raw = self.selection_text.toPlainText().strip()
        if not raw and not self.image_path:
            QMessageBox.warning(self, "Missing image", "Select an image first.")
            return

        if self.chk_use_vision_for_accuracy.isChecked():
            self._vision_reread_to_improved_text()
            return

        if not raw:
            QMessageBox.warning(self, "No selection OCR", "Drag-select an area first (OCR will auto-detect).")
            return

        fixed = self._ai_fix_ocr_text(raw)
        if fixed is not None:
            self.improved_text.setPlainText(fixed)

    def increase_accuracy_ai_context(self):
        raw = self.ai_context.toPlainText().strip()
        if not raw:
            QMessageBox.warning(self, "No context", "Put OCR text in the AI context box first.")
            return
        fixed = self._ai_fix_ocr_text(raw)
        if fixed is not None:
            self.ai_context.setPlainText(fixed)

    def _ai_fix_ocr_text(self, raw_text: str):
        client = self._ensure_groq()
        if client is None:
            return None

        try:
            self.set_status("Increasing accuracy (AI fix)...")
            model = self.model_box.currentText()

            system = (
                "You fix OCR text from blurry question papers.\n"
                "Strict rules:\n"
                "1) Do NOT invent missing questions/answers.\n"
                "2) Only repair when the intended text is very likely.\n"
                "3) Preserve numbering, options (A/B/C/D), punctuation, and math symbols.\n"
                "4) If unclear, keep the fragment and mark unclear parts as [??].\n"
                "5) Output only the corrected text (no explanation).\n"
            )

            completion = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": f"OCR TEXT:\n{raw_text}\n\nReturn corrected OCR text."}
                ],
                temperature=0.2,
                max_tokens=1200,
            )
            fixed = completion.choices[0].message.content or ""
            self.set_status("AI fix done. Review [??] parts.")
            return fixed.strip()
        except Exception as e:
            self.set_status("Increase accuracy failed.")
            QMessageBox.critical(self, "Increase Accuracy error", str(e))
            return None

    def vision_reread_to_ai_context(self):
        if not self.image_path:
            QMessageBox.warning(self, "Missing image", "Select an image first.")
            return

        client = self._ensure_groq()
        if client is None:
            return

        try:
            self.set_status("Vision re-read...")
            img = Image.open(self.image_path)

            if self.chk_vision_use_selection.isChecked():
                box = self._selection_crop_box_original_pixels()
                if box:
                    img = img.crop(box)

            data_url = self._image_to_data_url(img)
            vision_model = "meta-llama/llama-4-scout-17b-16e-instruct"

            completion = client.chat.completions.create(
                model=vision_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "Transcribe ALL visible text accurately.\n"
                                    "Do NOT add new content.\n"
                                    "Preserve numbering and line breaks.\n"
                                    "If unsure, write [??].\n"
                                    "Output ONLY the transcribed text."
                                )
                            },
                            {"type": "image_url", "image_url": {"url": data_url}},
                        ],
                    }
                ],
                temperature=0.1,
                max_tokens=1400,
            )

            text = (completion.choices[0].message.content or "").strip()
            self.ai_context.setPlainText(text)
            self.set_status("Vision re-read done.")
        except Exception as e:
            self.set_status("Vision re-read failed.")
            QMessageBox.critical(self, "Vision Re-read error", str(e))

    def _vision_reread_to_improved_text(self):
        if not self.image_path:
            QMessageBox.warning(self, "Missing image", "Select an image first.")
            return

        client = self._ensure_groq()
        if client is None:
            return

        try:
            self.set_status("Increasing accuracy (Vision re-read)...")
            img = Image.open(self.image_path)
            box = self._selection_crop_box_original_pixels()
            if box:
                img = img.crop(box)

            data_url = self._image_to_data_url(img)
            vision_model = "meta-llama/llama-4-scout-17b-16e-instruct"

            completion = client.chat.completions.create(
                model=vision_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "Transcribe ALL visible text accurately.\n"
                                    "Do NOT add new content.\n"
                                    "Preserve numbering and line breaks.\n"
                                    "If unsure, write [??].\n"
                                    "Output ONLY the transcribed text."
                                )
                            },
                            {"type": "image_url", "image_url": {"url": data_url}},
                        ],
                    }
                ],
                temperature=0.1,
                max_tokens=1400,
            )

            text = (completion.choices[0].message.content or "").strip()
            self.improved_text.setPlainText(text)
            self.set_status("Vision re-read done. Review [??] parts.")
        except Exception as e:
            self.set_status("Vision re-read failed.")
            QMessageBox.critical(self, "Vision Re-read error", str(e))

    def analyze_image_vision(self):
        if not self.image_path:
            QMessageBox.warning(self, "Missing image", "Select an image first.")
            return

        client = self._ensure_groq()
        if client is None:
            return

        try:
            self.set_status("Analyzing image with vision...")
            img = Image.open(self.image_path)

            if self.chk_vision_use_selection.isChecked():
                box = self._selection_crop_box_original_pixels()
                if box:
                    img = img.crop(box)

            data_url = self._image_to_data_url(img)
            vision_model = "meta-llama/llama-4-scout-17b-16e-instruct"

            completion = client.chat.completions.create(
                model=vision_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Describe what this image shows. Be structured."},
                            {"type": "image_url", "image_url": {"url": data_url}},
                        ],
                    }
                ],
                temperature=0.4,
                max_tokens=900,
            )

            text = completion.choices[0].message.content
            self.image_analysis.setPlainText((text or "").strip())
            self.set_status("Vision analysis done.")
        except Exception as e:
            self.set_status("Vision analysis failed.")
            QMessageBox.critical(self, "Vision error", str(e))


def apply_windows_like_style(app: QApplication):
    for style_name in ("windowsvista", "Windows", "Fusion"):
        try:
            app.setStyle(style_name)
            break
        except Exception:
            pass


if __name__ == "__main__":
    app = QApplication([])
    apply_windows_like_style(app)
    QApplication.setWindowIcon(QIcon(r"C:\Users\raidenxr\Downloads\Analyzer.ico"))


    w = Image2ExcelGroqApp()
    w.show()
    app.exec()
